# ============================================================
# modules/nlp_processor.py
# This file handles all NLP (Natural Language Processing) tasks:
# - Extracting text from PDF/DOCX files
# - Finding skills in text
# - Calculating how similar two texts are
# ============================================================

# PyPDF2 reads text from PDF files
import PyPDF2

# python-docx reads text from DOCX files
import docx

# re = regular expressions, for finding patterns in text
import re

# os for file operations
import os

# NLTK = Natural Language Toolkit (NLP library)
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Sklearn for TF-IDF and cosine similarity
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Download required NLTK data (only needed once)
# These are datasets that NLTK uses
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

try:
    nltk.download('punkt_tab', quiet=True)
except:
    pass


# ============================================================
# COMPREHENSIVE SKILLS DATABASE
# This is a large list of tech skills we will look for
# in resumes and job descriptions
# ============================================================

TECH_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'c++', 'c#', 'c', 'ruby', 'php', 'swift',
    'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'perl', 'typescript',
    'dart', 'julia', 'haskell', 'lua',

    # Web Development
    'html', 'css', 'react', 'angular', 'vue', 'node', 'nodejs', 'express',
    'django', 'flask', 'fastapi', 'spring', 'laravel', 'bootstrap', 'jquery',
    'next.js', 'nextjs', 'gatsby', 'svelte', 'redux', 'graphql', 'rest api',
    'restful', 'ajax', 'webpack', 'sass', 'scss', 'tailwind',

    # Data Science & ML
    'machine learning', 'deep learning', 'neural networks', 'tensorflow',
    'pytorch', 'keras', 'scikit-learn', 'sklearn', 'pandas', 'numpy', 'scipy',
    'matplotlib', 'seaborn', 'plotly', 'data science', 'data analysis',
    'data visualization', 'statistical analysis', 'feature engineering',
    'model training', 'model deployment', 'computer vision', 'opencv',
    'nlp', 'natural language processing', 'bert', 'transformers', 'gpt',
    'llm', 'hugging face', 'xgboost', 'random forest', 'regression',
    'classification', 'clustering', 'recommendation systems',

    # Databases
    'sql', 'mysql', 'postgresql', 'mongodb', 'sqlite', 'oracle', 'redis',
    'elasticsearch', 'cassandra', 'dynamodb', 'firebase', 'nosql',
    'database design', 'stored procedures',

    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins',
    'ci/cd', 'devops', 'terraform', 'ansible', 'linux', 'unix', 'bash',
    'shell scripting', 'git', 'github', 'gitlab', 'bitbucket', 'nginx',
    'apache', 'microservices', 'serverless',

    # Data Tools
    'tableau', 'power bi', 'excel', 'hadoop', 'spark', 'kafka', 'airflow',
    'jupyter', 'databricks', 'snowflake', 'dbt', 'etl', 'data pipeline',
    'data warehouse', 'data lake',

    # Mobile
    'android', 'ios', 'react native', 'flutter', 'xamarin',

    # Security
    'cybersecurity', 'network security', 'penetration testing', 'ethical hacking',
    'cryptography', 'firewalls',

    # Soft Skills (also important!)
    'communication', 'teamwork', 'leadership', 'problem solving', 'agile',
    'scrum', 'project management', 'critical thinking', 'time management',

    # Other Technical
    'api', 'json', 'xml', 'microservices', 'oop', 'object oriented',
    'data structures', 'algorithms', 'design patterns', 'unit testing',
    'test driven', 'tdd', 'version control', 'agile', 'scrum'
}


# ============================================================
# FUNCTION 1: Extract text from uploaded file
# ============================================================

def extract_text_from_file(filepath):
    """
    Read a PDF or DOCX file and return all the text inside it.

    filepath = the path to the file on the server
    Returns a string of all the text.
    """
    # Get the file extension (.pdf or .docx)
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(filepath)
    else:
        return ""


def extract_text_from_pdf(filepath):
    """Read all text from a PDF file"""
    text = ""
    try:
        # Open the PDF file in binary read mode ('rb')
        with open(filepath, 'rb') as file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(file)

            # Loop through every page and extract text
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:  # Only add if there's actual text
                    text += page_text + "\n"

    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

    return text.strip()  # Remove extra whitespace


def extract_text_from_docx(filepath):
    """Read all text from a DOCX file"""
    text = ""
    try:
        # Open the DOCX document
        doc = docx.Document(filepath)

        # Loop through all paragraphs and add their text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text += paragraph.text + "\n"

        # Also check tables inside the DOCX
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

    return text.strip()


# ============================================================
# FUNCTION 2: Extract skills from text
# ============================================================

def extract_skills(text):
    """
    Find all tech skills mentioned in a piece of text.

    text = a string (resume text or job description)
    Returns a list of skills found.
    """
    if not text:
        return []

    # Convert text to lowercase for comparison
    text_lower = text.lower()

    found_skills = set()  # Use a set to avoid duplicates

    # Check each skill in our database
    for skill in TECH_SKILLS:
        # Use word boundary matching to find exact skills
        # e.g., "python" should not match "pythonic"
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.add(skill)

    return list(found_skills)


# ============================================================
# FUNCTION 3: Calculate TF-IDF Cosine Similarity
# ============================================================

def calculate_similarity(resume_text, job_description):
    """
    Calculate how similar the resume is to the job description.

    This uses:
    - TF-IDF: Converts text to numbers (vectors)
      TF = Term Frequency (how often a word appears)
      IDF = Inverse Document Frequency (how rare is the word)
    - Cosine Similarity: Measures the angle between two vectors
      Score of 1.0 = identical, Score of 0.0 = completely different

    Returns a score from 0 to 100.
    """
    if not resume_text or not job_description:
        return 0.0

    # Clean the texts
    resume_clean = clean_text(resume_text)
    jd_clean     = clean_text(job_description)

    try:
        # Create TF-IDF Vectorizer
        # stop_words='english' removes common words like 'the', 'and', 'is'
        # These don't carry meaning and would skew the score
        vectorizer = TfidfVectorizer(
            stop_words='english',    # Remove stopwords
            ngram_range=(1, 2),      # Consider 1 and 2 word phrases
            min_df=1,                # Minimum document frequency
            max_features=1000        # Use top 1000 features
        )

        # Fit and transform: convert texts to TF-IDF vectors (numbers)
        # We pass both texts together so vectorizer knows all words
        tfidf_matrix = vectorizer.fit_transform([resume_clean, jd_clean])

        # Calculate cosine similarity between the two vectors
        # tfidf_matrix[0] = resume vector, tfidf_matrix[1] = JD vector
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])

        # similarity is a 2D array like [[0.75]]
        # We extract the number and convert to 0-100 scale
        score = float(similarity[0][0]) * 100

        return round(score, 2)

    except Exception as e:
        print(f"Error calculating similarity: {e}")
        return 0.0


def clean_text(text):
    """
    Clean text for better analysis:
    - Remove special characters
    - Convert to lowercase
    - Remove extra spaces
    """
    # Remove URLs
    text = re.sub(r'http\S+|www\S+', '', text)

    # Remove email addresses
    text = re.sub(r'\S+@\S+', '', text)

    # Remove phone numbers
    text = re.sub(r'[\+\d]?(\d{2,3}[-\.\s]??\d{2,3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4})', '', text)

    # Remove special characters but keep spaces and alphanumeric
    text = re.sub(r'[^\w\s]', ' ', text)

    # Convert to lowercase
    text = text.lower()

    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


# ============================================================
# FUNCTION 4: Get keyword frequencies for visualization
# ============================================================

def get_keyword_frequencies(text, top_n=20):
    """
    Count how many times each important word appears.
    Used for the word frequency chart in the frontend.

    Returns a dict like: {'python': 5, 'data': 8, ...}
    """
    if not text:
        return {}

    # Clean and tokenize
    clean = clean_text(text)

    try:
        # Get English stopwords (common words to ignore)
        stop_words = set(stopwords.words('english'))
        # Add our own custom stopwords
        stop_words.update(['experience', 'work', 'using', 'use', 'years', 'year',
                           'strong', 'good', 'knowledge', 'ability', 'skills'])

        # Split text into individual words
        words = word_tokenize(clean)

        # Count frequencies, ignoring stopwords and short words
        freq = {}
        for word in words:
            if word not in stop_words and len(word) > 3:
                freq[word] = freq.get(word, 0) + 1

        # Sort by frequency and return top N
        sorted_freq = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return dict(sorted_freq[:top_n])

    except Exception as e:
        print(f"Error getting frequencies: {e}")
        # Fallback: simple word count
        words = clean.split()
        freq  = {}
        for w in words:
            if len(w) > 3:
                freq[w] = freq.get(w, 0) + 1
        sorted_freq = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return dict(sorted_freq[:top_n])
